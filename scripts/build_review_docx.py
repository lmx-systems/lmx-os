"""
Build the cofounder review as a .docx in house style (docs/DOCUMENT_STYLE.md).

Why a script rather than a one-off export: this document will be revised after the
review, and re-applying Aptos, the 0.5" margins, the brand green and the logo by hand
each time is exactly the retrofitting the house style exists to avoid.

Reads the published narrative's structure from a plain-text source of truth kept in
this file, so the Word version cannot drift from the HTML silently - if the narrative
changes, this changes with it in the same commit.

Usage:
    .venv/bin/python -m scripts.build_review_docx
    .venv/bin/python -m scripts.build_review_docx --out ~/Desktop/review.docx

Output is deliberately NOT committed - `.gitignore` covers *.docx and the shared drive
is the source of truth for business documents.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BRAND = RGBColor(0x0A, 0x66, 0x44)
BRAND_HEX = "0A6644"
TINT_HEX = "E6F1EB"
INK = RGBColor(0x14, 0x17, 0x1C)
INK_SOFT = RGBColor(0x3A, 0x41, 0x50)
INK_MUTED = RGBColor(0x5B, 0x64, 0x72)

REPO = Path(__file__).resolve().parent.parent
LOGO = REPO / "docs" / "LMX branding " / "lmx-logo-lockup-light.png"

FONT = "Aptos"


def _shade(cell, hex_fill: str) -> None:
    """Table cell background. python-docx has no API for this, so it goes in as XML."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _rule(paragraph, hex_color: str = BRAND_HEX, size: int = 12) -> None:
    """A coloured rule under a paragraph - used for the masthead and section heads."""
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def _setup(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.22
    # East-Asian font name too, or Word substitutes for any non-Latin run.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def _logo(doc: Document) -> None:
    """The lockup at the top of page one, per house style - not a text label."""
    if not LOGO.exists():
        print(f"warning: logo not found at {LOGO} - continuing without it", file=sys.stderr)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(LOGO), height=Inches(0.30))


def _kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = BRAND
    run.font.name = FONT


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = INK
    run.font.name = FONT


def _byline(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = INK_MUTED
    run.font.name = FONT
    _rule(p, BRAND_HEX, size=16)


def _heading(doc: Document, number: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    n = p.add_run(f"{number}   ")
    n.font.size = Pt(10)
    n.font.bold = True
    n.font.color.rgb = BRAND
    n.font.name = FONT
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = INK
    run.font.name = FONT
    _rule(p, "D8DCE2", size=6)


def _body(doc: Document, text: str, *, lead: bool = False, indent: float = 0.0) -> None:
    """One paragraph. **bold** spans are honoured; nothing else is parsed - the
    narrative is prose, and a full markdown renderer here would be more surface than
    the document needs."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = p.add_run(chunk)
        run.font.name = FONT
        run.font.size = Pt(11.5 if lead else 10.5)
        run.font.bold = i % 2 == 1
        run.font.color.rgb = INK_SOFT if lead else INK


def _quote(doc: Document, text: str) -> None:
    """A pull quote, marked with a left brand rule rather than italics - it is an
    assertion we want read, not an aside."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BRAND_HEX)
    bdr.append(left)
    p._p.get_or_add_pPr().append(bdr)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def _sub(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = INK


def _screen(doc: Document, where: str, lines: list[str]) -> None:
    """A screen rendered as a transcript rather than a picture.

    The web version draws the frames. Word cannot, and a screenshot pasted in would
    be unreadable at this width - so the screen goes in as what it actually says, in
    order, which is the part being reviewed anyway. One shaded cell so it reads as a
    quoted object rather than as more prose."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    _shade(cell, TINT_HEX)
    cell.text = ""

    head = cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(5)
    run = head.add_run(where.upper())
    run.font.name = "Aptos Mono"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = BRAND

    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        # A leading tab marks a field's value or hint - one level of indent, so the
        # label/value relationship survives without a nested table.
        depth = len(line) - len(line.lstrip("\t"))
        if depth:
            p.paragraph_format.left_indent = Inches(0.18 * depth)
        for i, chunk in enumerate(line.strip().split("**")):
            if not chunk:
                continue
            run = p.add_run(chunk)
            run.font.name = FONT
            run.font.size = Pt(9.5 if depth == 0 else 9)
            run.font.bold = i % 2 == 1
            run.font.color.rgb = INK if depth == 0 else INK_MUTED
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _note(doc: Document, number: int, claim: str, why: str) -> None:
    """One numbered design decision: the claim in bold, the reasoning under it."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.32)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    tag = p.add_run(f"{number:>2}  ")
    tag.font.name = "Aptos Mono"
    tag.font.size = Pt(9)
    tag.font.bold = True
    tag.font.color.rgb = BRAND
    run = p.add_run(claim)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = INK

    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Inches(0.32)
    q.paragraph_format.space_after = Pt(2)
    run = q.add_run(why)
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.color.rgb = INK_MUTED


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _shade(cell, BRAND_HEX)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header.upper())
        run.font.name = FONT
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if index % 2 == 1:
                _shade(cells[i], TINT_HEX)
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(value)
            run.font.name = FONT
            run.font.size = Pt(9)
            run.font.bold = i == 0
            run.font.color.rgb = INK if i == 0 else INK_SOFT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# The document
# ---------------------------------------------------------------------------

TENETS = [
    ("One canonical order object. Many doors. A core that knows nothing.",
     "If a new way of receiving orders requires a change inside routing, batching or the driver app, the contract is wrong and we fix the contract.",
     "Settles: whether to special-case a customer. No."),
    ("Telling them counts as much as taking the order.",
     "A carrier that accepts work and goes quiet is a favour, not a carrier. Write-back is an exit criterion for every piece of work, never a follow-up ticket.",
     "Settles: what gets cut when a deadline slips. Not this."),
    ("A claim is not a fact.",
     "Where something matters - a licence expiry, a delivery time, proof that parts arrived - the system records who established it and when. It never presents something a person asserted as something we checked.",
     "Settles: whether a convenient default is acceptable. Not if it fabricates a check."),
    ("Nothing is dispatched unpriced, and nothing silently free.",
     "A client is approved and priced in the same action. An order we cannot price is refused loudly rather than delivered at zero.",
     "Settles: whether to onboard someone quickly and sort rates later. No."),
    ("Refuse rather than guess.",
     "An address we cannot resolve, a distance we cannot compute, a proof we did not capture - each is a visible refusal. A wrong coordinate sends a real van to the wrong place; no coordinate merely stops an order.",
     "Settles: whether to fall back to a default. Only when the default cannot be mistaken for truth."),
    ("The driver's screen shows one decision.",
     "No tiers, no deadlines, no negotiating room. If a choice can be made by the system with better information, it is not the driver's to make at a doorstep.",
     "Settles: whether to surface something 'just in case'. No."),
    ("We are an operator, not a software vendor.",
     "Nothing here is licensed or sold. We are paid per delivery. Everything a customer touches is a way of handing us work, and it stays unnamed as a product.",
     "Settles: any conversation that starts with 'could we sell this to...'. No."),
]

FAQ = [
    ("Who is this actually for - contracted overflow, or walk-ups?",
     "This is the open question with the widest consequences, and it is a commercial one rather than a technical one. If we are serving distributors we have contracts with and invoice monthly, we need no payment collection at all and the front door is a convenience for their staff. If we are serving genuine walk-ups, we need to take money before we carry anything, which is a month of work and a compliance conversation. Building it speculatively wastes the month; discovering we need it mid-pilot is worse."),
    ("If a distributor can sign up on a link, what stops anyone signing up?",
     "Nothing stops them applying, and everything stops them dispatching. An application creates a pending account whose users cannot log in. A person at LMX approves it, and sets their prices in the same action. Until then it is a lead, not a customer."),
    ("We are giving a customer's customer a live map of one of our vans. Is that wise?",
     "Only under a rule we should be explicit about. The driver appears on that map only while they are on their way to that specific address - never while they are delivering somebody else's parcel, which would tell one recipient roughly where another lives and show both of them the shape of the driver's working day. The destination is hinted rather than printed in full, because these links get forwarded. And the link stops working a day after the delivery, so it does not become a permanent window onto whoever is carrying that route next week."),
    ("What happens when we are late?",
     "A credit appears on that month's statement, itemised against the specific order, with how many minutes late it was. The distributor does not have to notice and ask. That is a deliberate choice to be held to a number rather than to a conversation - and it means the service level we agree to has to be one we can actually hit, because it now costs us money when we don't."),
    ("What stops a driver quietly marking a delivery done that never happened?",
     "A completion has to carry the proof that order asked for - the photo count, the named subjects, the signature, or a PIN we texted the customer and verified against what we issued. Where cash is owed, the stop cannot be closed with the money unaccounted for: either it was collected in full, or a dispute went to the distributor. And every collection names the driver who took it, because cash in a van is a custody question."),
    ("What is the one thing most likely to eat a week during onboarding?",
     "The point-of-sale field names. Our connector for the most common system was written against documentation rather than a live tenant, and every one of those integrations differs slightly in what it calls a ship-to address. It should be checked against a real system on day one of the first engagement rather than discovered in week two."),
    ("How do we know any of this is working, rather than believing it?",
     "Four numbers, computed from real records rather than asserted: how long from approving a customer to their first delivery landing, how long a counter person takes to enter an order from their second one onward, how quickly a status change reaches them, and how many orders needed a human at LMX to correct them. The fourth one we cannot measure yet and the system says so rather than reporting a zero."),
    ("What is the biggest thing this does not yet handle?",
     "A partial delivery. Three of four parts is an ordinary situation in this trade and there is currently no way to record it - a driver's only options are delivered or failed. It needs a decision about what gets billed before it can be built, which makes it a question for this review rather than a task."),
]

# The six screens, carried over and extended from the 7 August annotated review.
# Notes are numbered continuously across all six because they are cited in argument
# ("note 14") and a per-screen restart would make two notes share a number.
SCREENS = [
    {
        "title": "1  ·  The page we send a prospect",
        "caption": "Public, no login, one URL. Submitting it applies for an account. It does not grant one.",
        "where": "lmx / signup",
        "lines": [
            "**Send deliveries with LMX**",
            "\tTell us about your business - takes a minute",
            "Company name",
            "\tMidtown Auto Parts",
            "Your name",
            "\tJordan Rivera",
            "Email",
            "\tjordan@midtownparts.com",
            "Phone",
            "\tOptional",
            "Where do you deliver?",
            "\tAustin metro",
            "\tRoughly where your deliveries go - a city or area is fine.",
            "Choose a password",
            "\tAt least 10 characters. You'll use this once your account is approved.",
            "[ ]  I agree to LMX's terms of service and privacy policy.",
            "**[ Request an account ]**",
            "\tAlready have an account? Sign in",
        ],
        "notes": [
            ("Six fields, every one answerable from memory.",
             "Nothing that needs a contract, a rate negotiation, or their IT department. This is the front "
             "of the funnel, and it is exactly where one extra required field costs us applicants."),
            ("\"Where do you deliver?\" is free text, not a dropdown.",
             "We have no service-area model, so nothing could route an applicant automatically. A person "
             "reads this and picks the hub at approval. That is honest, and better than a dropdown that "
             "pretends to know."),
            ("They choose their own password before we approve them.",
             "So approving is one click, with no credential to mint and email. The account cannot do "
             "anything until approval, so a password existing early grants nothing."),
            ("Submitting says \"we'll be in touch\" and nothing else.",
             "No reference, no next step - and an email already on our books gets the identical response. "
             "Otherwise this page becomes a way for anyone to test who our customers are."),
        ],
        "ask": "the checkbox names a document. Whatever version an applicant accepts is recorded against "
               "them, which means the sentence has to be true on the day the page goes live. That is the "
               "first item in section 9 and it gates this screen specifically.",
    },
    {
        "title": "2  ·  What we do with the application",
        "caption": "Our side. An application arrives here and waits. Nothing about it can move a van yet.",
        "where": "lmx ops / signups",
        "lines": [
            "**Midtown Auto Parts**                              PENDING",
            "\tJordan Rivera  ·  jordan@midtownparts.com",
            "\tDelivers around: Austin metro  ·  Applied 2 hours ago",
            "Hub",
            "\tAustin - Braker",
            "What we charge them, per tier",
            "\tHot shot $28.00      Urgent $18.00",
            "\tStandard $12.00      Scheduled $9.50",
            "**[ Approve and set prices ]**    [ Reject ]",
        ],
        "notes": [
            ("Approval and pricing are the same action.",
             "You cannot activate a customer without saying what they pay. That single constraint removes "
             "an entire class of problem downstream: a live customer with no price on file, and an order "
             "we cannot bill."),
            ("A rejected applicant and a former customer are different states.",
             "Tempting to reuse one flag for both. They behave differently - one never had access, one had "
             "it and lost it - and collapsing them would eventually let the wrong one back in."),
            ("A person picks the hub.",
             "The deliberate manual step from note 2. It is the one place where our lack of a service-area "
             "model surfaces as work, and it is a minute per customer."),
        ],
        "ask": "these four prices are the whole commercial relationship, typed once at approval. Do we want "
               "an approver to be able to set them freely, or should they start from a house rate card that "
               "has to be deliberately overridden?",
    },
    {
        "title": "3  ·  Sending a delivery",
        "caption": "The screen a counter person uses with a customer on the phone. Under a minute for the "
                   "first order, under thirty seconds for the ones after it.",
        "where": "lmx portal / new order",
        "lines": [
            "Deliver to",
            "\t900 Congress Ave, Austin TX",
            "Collect from",
            "\t( 1200 E 6th St )   ( Riverside Depot )   ( Braker Ln )",
            "\tWe'll remember it - next time it's one tap.",
            "When",
            "\tNow - straight there, no waiting        **Within the hour - urgent**",
            "\tToday - standard                        Tomorrow - scheduled",
            "Add contact, notes or a reference",
            "**[ Send this delivery ]**",
        ],
        "notes": [
            ("Three fields on the fast path.",
             "Where it goes, where it's from, how soon. Contact, notes and reference all exist, collapsed "
             "behind one line, so they never slow down the order that doesn't need them."),
            ("Destination first, pickup second.",
             "Backwards from how we think about it and correct for how they work. The destination is what "
             "the person on the phone is reading out right now. The pickup is usually their own shop and "
             "usually one tap."),
            ("Previous pickups are chips, not a dropdown.",
             "A distributor collects from the same handful of places forever. Type an address once, it "
             "becomes a saved location, and the next order to it is a tap. That is where most of the "
             "thirty seconds comes from."),
            ("Four choices instead of a date picker.",
             "Nobody at a counter operates a calendar widget. These map onto our service tiers behind the "
             "scenes: the client says how urgent it is, we decide what that obliges us to."),
            ("This form is the smallest of four doors.",
             "Pasting a list, uploading a spreadsheet, or their own system calling ours all arrive as the "
             "same order. The form is what a counter uses all day; the others are back-office and machine "
             "traffic. Deliberately, none of them is privileged."),
        ],
        "ask": "the four urgency choices are also, implicitly, four prices. A counter person picking \"Now\" "
               "out of habit is choosing our most expensive tier. Should the price of the choice be on the "
               "button?",
    },
    {
        "title": "4  ·  The confirmation",
        "caption": "The screen with the most riding on it. This is what makes us read as a carrier rather "
                   "than as a form that swallowed something.",
        "where": "lmx portal / booked",
        "lines": [
            "**Booked - we'll collect by 2:40 PM**",
            "\tEstimated delivery around 3:25 PM",
            "\tReference   LMX-7C2A9F",
            "\tPrice       $18.00",
            "**[ Send another ]**",
        ],
        "notes": [
            ("A commitment, not a spinner.",
             "\"We'll collect by 2:40\" is the whole difference between a carrier and a contact form. The "
             "price sits beside it because a counter person is frequently quoting their customer while "
             "they type."),
            ("The two times are deliberately unequal.",
             "Collection is a promise, derived from the service level they bought, and missing it credits "
             "their statement automatically. Delivery is an estimate, and it is set quieter and worded as "
             "one. Nothing on this screen invites a customer to read the second number as the first."),
            ("\"Send another\" goes straight to an empty form.",
             "The realistic pattern is three orders in a row, not one and done. Anything that makes the "
             "second order slower than the first is the wrong shape."),
        ],
        "ask": "we show the delivery estimate because warmth is worth something at the counter. The cost of "
               "being wrong is a customer who read it as a promise. The wording carries that distinction "
               "today - is wording enough?",
    },
    {
        "title": "5  ·  The driver's stop",
        "caption": "One screen, in a van, one-handed, sometimes with no signal. The driver never chooses "
                   "what to do next - the route does.",
        "where": "driver app / stop 3 of 7",
        "lines": [
            "**Drop  ·  Stop 3 of 7**                                ARRIVED",
            "\t900 Congress Ave, Suite 400",
            "\tAsk for the service desk",
            "**Proof needed here**",
            "\t2 photos - the parcel at the door",
            "\tPlus a signature or a PIN",
            "\t( Photo 1 done )   ( Photo 2 )   ( Signature )",
            "**Collect $142.60 on delivery**",
            "\tCash or check, in full",
            "\t( Cash )   ( Check )   ( They won't pay )",
            "**[ Complete stop ]**",
        ],
        "notes": [
            ("\"Stop 3 of 7\" is the entire navigation.",
             "No list to browse, no stop to pick. Work is pushed, not chosen. It removes the cherry-picking "
             "problem before it exists, and it means a new driver's first shift needs no explanation."),
            ("What counts as proof is set per order, and the screen says so.",
             "A brake rotor left at a loading dock and a $4,000 ECU handed to a service manager are not the "
             "same delivery. The customer decides which one they bought; the driver just reads it. Nothing "
             "completes until the evidence the order asked for actually exists."),
            ("Cash on delivery has no amount field.",
             "The driver confirms the amount we already know, or says the customer refused, which raises it "
             "to a person immediately. A driver typing what they collected turns every shortfall into a "
             "dispute about arithmetic instead of a dispute about the customer."),
            ("Every action here works with no signal.",
             "Photos, signature, completion, cash - all queue locally and sync when the phone reconnects, "
             "and the driver can see plainly what has gone up and what has not. A parking garage is not an "
             "outage."),
        ],
        "ask": "the driver cannot record a partial delivery - four of six pieces arriving is either a "
               "completion or an exception. Making it real means deciding what we bill for it, which is a "
               "commercial question before it is a screen.",
    },
    {
        "title": "6  ·  What the person waiting for the part sees",
        "caption": "A texted link, no login, no app. Our customer's customer - the only screen here that "
                   "someone who has never heard of us will look at.",
        "where": "track.lmx / recipient",
        "lines": [
            "**Your delivery from Midtown Auto Parts**",
            "\tArriving on Congress Ave",
            "**On the way to you**                               NEXT STOP",
            "\tExpected around 3:25 PM",
            "\t[ map - the van, shown only while this drop is next ]",
            "\tCollected 2:38 PM  ·  Midtown Auto Parts",
            "\tQuestions about this delivery? Contact Midtown Auto Parts",
        ],
        "notes": [
            ("The van appears only when they are genuinely next.",
             "A moving dot that is four stops away is worse than no dot - it invites someone to stand "
             "outside. Position is shown when it means something and withheld when it doesn't."),
            ("The link stops working shortly after delivery.",
             "A tracking URL that lives forever is a permanent window into someone's address and order "
             "history. It expires, and an expired link and a link that never existed are indistinguishable "
             "from the outside."),
            ("Our customer's name is on it, not just ours.",
             "We are the carrier; the relationship is theirs. Questions route back to the distributor. This "
             "screen is a reason for them to keep sending us work, not a way for us to meet their customers."),
        ],
        "ask": "this is the most widely seen screen we have and the only one an outsider judges us on. It is "
               "also the one with no branding decision behind it yet.",
    },
]

CAPABILITIES = [
    ["Self-serve signup", "A shareable, embeddable link. Application creates nothing that can act; approval sets prices in the same step."],
    ["Four ways to send an order", "A form, a paste, a CSV upload, or their own system over an API key. All become the same order."],
    ["Address memory", "A typed pickup becomes a saved location on first use. The second order to it needs no typing."],
    ["Urgency without jargon", "Four plain choices at the counter; the system classifies the tier and commits to a collection time."],
    ["Batching", "Orders are held briefly so nearby ones combine into a single trip. The economic engine."],
    ["Fleet-wide routing", "Re-optimised on every meaningful event across all open orders and all available drivers."],
    ["Offer-based dispatch", "A driver accepts or declines; a decline returns the work to the pool."],
    ["Configurable proof", "Per order: how many photos, of what, plus a signature or a verified PIN."],
    ["Cash on delivery", "Collect in full or escalate. No partial amounts, by design."],
    ["Driver compliance gate", "Licence and insurance, uploaded and verified by a person, before anyone goes on shift."],
    ["Offline operation", "Every driver action works with no signal and syncs later, with honest sync state."],
    ["Customer tracking", "A texted link showing status, an ETA, and the van when it is genuinely inbound. Expires after delivery."],
    ["Status write-back", "To the portal, to the recipient, and to their own systems as signed messages retried for three days."],
    ["Rate tables", "Base plus per mile, per piece, per weight, with a minimum. Every fee records the arithmetic behind it."],
    ["Automatic credits", "A missed commitment credits the statement, itemised, without being asked."],
    ["Returns and cores", "Collected on the delivery visit, expected or ad hoc."],
    ["Exception handling", "Five reasons a driver can flag, each routed to whoever can resolve it."],
    ["Self-measurement", "Onboarding time, order entry time, write-back latency - from records, not impressions."],
]


def build(out: Path) -> None:
    doc = Document()
    _setup(doc)

    _logo(doc)
    _kicker(doc, "Product & design review · for cofounder discussion")
    _title(doc, "How a distributor sends us orders")
    _byline(doc, "A six-pager, written to be read start to finish  ·  11 August 2026  ·  Sourabh Miglani")

    _heading(doc, "1", "The whole thing in one paragraph")
    _body(doc, "A parts distributor decides on a Tuesday morning that they want us delivering for them. "
               "By Tuesday afternoon their counter staff are sending us jobs, their customers are getting "
               "live tracking links, and their accounts team knows exactly what it will cost. Nobody at "
               "their company wrote any software, exported any data, or spoke to anyone in ours. That is "
               "the entire product: **we absorb whatever they already have - a person at a counter, a "
               "spreadsheet, a CSV export, their own point-of-sale system - and we send status back "
               "without being asked.**", lead=True)

    _heading(doc, "2", "The press release we would write")
    _body(doc, "Written first, and backwards from the customer, because if this paragraph is not worth "
               "reading then nothing behind it is worth building.")
    _sub(doc, "LMX starts same-day parts delivery for distributors who don't want an IT project")
    _body(doc, "Counter staff send a delivery in under thirty seconds from any phone or browser. Their "
               "customers watch the van arrive. No integration and nothing to install - LMX is paid per "
               "delivery.", indent=0.2)
    _body(doc, "**AUSTIN, TEXAS -** LMX today began accepting same-day delivery orders from parts "
               "distributors with no setup work required on the distributor's side. A branch manager can "
               "sign up on a link, be approved the same morning, and have their first delivery collected "
               "that afternoon.", indent=0.2)
    _body(doc, "Same-day parts delivery has historically required either a dedicated courier relationship "
               "negotiated per branch, or a software platform the distributor has to configure, license "
               "and train staff on. Both take weeks. Both assume the distributor has someone whose job is "
               "logistics software. Most branches do not: they have a counter, a phone, and a person who "
               "already has a queue of customers in front of them.", indent=0.2)
    _body(doc, "LMX takes orders in whatever form the branch already produces them. A counter person "
               "types one delivery in under thirty seconds. A dispatcher pastes twenty lines out of a "
               "spreadsheet. An office manager uploads the CSV their system already exports. A "
               "distributor with their own software posts orders to us directly and receives status back "
               "automatically. Every one of those arrives as the same order, and every one is dispatched "
               "by the same routing system.", indent=0.2)
    _body(doc, "Status comes back without anyone chasing it. The distributor watches each order move. "
               "Their customer gets a text with a live tracking link when the parts are collected, showing "
               "the van approaching. If a delivery misses its promised window, the credit appears on the "
               "statement automatically, itemised, without the distributor having to notice and ask.",
          indent=0.2)
    _body(doc, "“We are not selling anyone software,” said Sourabh Miglani, CTO. “We are a "
               "delivery company that made itself easy to hand work to. The measure of that is whether a "
               "branch that said yes this morning is sending us orders this afternoon - and whether they "
               "can tell, without calling us, exactly where every one of them is.”", indent=0.2)

    _heading(doc, "3", "What it is for the distributor")
    _body(doc, "Start at the counter, because that is where the order comes from and the counter person is "
               "the hardest user to win. They are mid-conversation with a customer who wants a rotor "
               "today. They have a queue. They will use a system that is faster than picking up the phone "
               "and abandon anything that is not.")
    _body(doc, "So the order form asks for the delivery address first, because that is the thing they have "
               "in their hand. It remembers every place they have ever collected from, so the second order "
               "to a shop is two taps rather than an address typed again. It asks when as a choice between "
               "four plain words - now, within the hour, today, tomorrow - because a counter person does "
               "not know what a service tier is and should never have to learn. It never blocks on a field "
               "they have not filled in.")
    _body(doc, "When they submit, they do not get a spinner and a confirmation number. They get a "
               "commitment: we will collect by 1:25, estimated delivery around 2:05, eighteen dollars. "
               "That is what they need in order to turn back to the customer in front of them and say "
               "something true.")
    _quote(doc, "The test of the counter experience is not whether it is elegant. It is whether the second "
                "order of the day takes under thirty seconds, and we measure that on every order rather "
                "than asserting it.")
    _body(doc, "Behind the counter, the same system serves three other people at the same distributor. A "
               "dispatcher with a day's worth of deliveries pastes them in as lines, or uploads the file "
               "their counter system exports, and gets a row-by-row account of what we understood - one "
               "unreadable address never discards the other thirty-nine. An office manager sees statements "
               "that show what was charged, what was credited back when we missed, and which specific "
               "orders those credits were for, with how late each one was. An owner adds and removes their "
               "own staff without contacting us.")
    _body(doc, "And a distributor who does have their own software gets an API key and a documented order "
               "format written in their vocabulary, not ours. Retries are safe, so a timeout never becomes "
               "a duplicate van. They can look any order up by their own reference rather than storing one "
               "of ours. Status comes back to them as signed messages their system can trust, retried for "
               "three days if their server is down, with a log they can debug from themselves.")

    _heading(doc, "4", "What it is for the driver")
    _body(doc, "A driver is not a user of this system so much as the mechanism by which it happens. They "
               "are in a van, in traffic, sometimes in the rain, holding a phone in one hand. Every design "
               "decision on their side follows from that.")
    _body(doc, "They log in with a phone number and a code - no password to remember or reset in a car "
               "park. Before they can go on shift, the system checks that their licence and insurance are "
               "on file and have been verified by someone at LMX, and if they are not, it says so plainly "
               "and tells them whose move it is next: upload this, or wait for us to check it. It never "
               "simply refuses.")
    _body(doc, "Work arrives as an offer with a countdown, not an assignment. They accept or they decline, "
               "and declining puts the orders back in the pool rather than stranding them. Once accepted, "
               "they see one ordered day, with the stop they are driving to marked as such.")
    _body(doc, "At each stop they get exactly what they need to complete it and nothing else: the address, "
               "who to ask for, the access notes about the rear dock and the bell on the left, how many "
               "parcels, and what proof this particular order requires - because one distributor wants a "
               "photo and another wants four with named subjects and a signature. They learn that on the "
               "way there, not at the door with the box already on the counter.")
    _body(doc, "Where money changes hands, the amount is on the screen before they knock, and there are "
               "exactly two things they can do: take all of it, or flag a dispute that goes straight to "
               "the distributor. **There is no field to type a smaller number into.** That is deliberate "
               "and it is for the driver's protection as much as ours: the money is the distributor's "
               "invoice to their own customer, nobody at LMX has the authority to discount it, and a "
               "driver arguing about it on a doorstep is arguing on somebody else's behalf.")
    _body(doc, "Three things they never see: a service tier, a deadline, or an urgency flag. The system "
               "manages all of that invisibly and hands them stops in the right order. Telling a driver "
               "that this one is a T1 would invite them to make a judgement the optimizer has already made "
               "with more information.")
    _body(doc, "And every tap works with no signal. Arriving, scanning, completing, flagging - all of it "
               "writes to the phone first and syncs when there is service, with a small honest indicator "
               "of what is still waiting. A driver in a basement car park is not blocked by us.")

    _heading(doc, "5", "What it is for LMX")
    _body(doc, "For the company, this is three things: a sales weapon, an operating system, and a source of "
               "compounding data.")
    _body(doc, "**As a sales weapon**, it removes the two objections that kill a courier conversation. "
               "“We'd need to involve IT” - no, you would not. “How would we know where our "
               "parts are?” - here is the link your customer gets. A branch manager can be delivering "
               "with us the same day they decide to, which means the sales cycle is a conversation rather "
               "than a project.")
    _body(doc, "**As an operating system**, it gives one orchestrator visibility of the whole hub on one "
               "screen: what is open, what is being deliberately held so nearby orders can be combined "
               "into one trip, who is out there and where. The holding is the economic engine - batching "
               "is what turns a set of single trips into a route, and it is the mechanism the "
               "deliveries-per-hour advantage rests on. Two review queues sit alongside it, because two "
               "things must never be automatic: approving a new client, which sets their prices in the "
               "same action so nothing can ever be delivered unpriced, and verifying a driver's documents, "
               "where the reviewer reads the expiry off the licence rather than accepting what the driver "
               "typed.")
    _body(doc, "**As a data asset**, every order carries its own ground truth: when it was received, when "
               "it was promised, when it was actually collected and delivered, what proof was captured, "
               "how long the person entering it took. That is what lets us answer questions about "
               "ourselves with numbers instead of impressions. Those are measured from real records rather "
               "than asserted, and where we cannot measure something honestly we say so rather than "
               "producing a number that looks like data.")
    _quote(doc, "The architectural bet worth understanding: one canonical order object, many doors in, many "
                "channels out, and a core that knows nothing about which door an order came through. It "
                "has now been tested twice - adding a client-facing API and a CSV importer required no "
                "change to batching, routing, or the driver app.")

    _heading(doc, "6", "Tenets")
    _body(doc, "These are ranked, and they exist to settle arguments rather than to inspire. When two of "
               "them conflict, the lower number wins.")
    for i, (name, text, settles) in enumerate(TENETS, start=1):
        _sub(doc, f"{i}. {name}")
        _body(doc, text, indent=0.2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(settles)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = INK_MUTED

    _heading(doc, "7", "What it deliberately does not do")
    _body(doc, "Scope discipline is easier to review than scope ambition, so here is what we have chosen "
               "against and why. Each of these is a decision, not an oversight.")
    for text in [
        "**It does not take card payments.** Cash at the door works, on the distributor's behalf. Everything else is a monthly invoice. Whether we ever need more than that depends on who we are actually serving, which is question one in the FAQ below.",
        "**It does not do EDI.** If a signed enterprise agreement requires it we will buy a broker connection. We will not build one.",
        "**It does not let a customer set their own service level.** They tell us when they need it; we classify the urgency and commit to a collection time. Otherwise the fastest tier is free to anyone who types an aggressive date.",
        "**It does not commingle across distributors on the premium tier.** A hot shot is point-to-point, and combining it with another shop's pickup is precisely the thing the premium exists to prevent.",
        "**It is not a portal we are trying to make people live in.** The best version of this is one a distributor stops noticing, because their own system talks to ours.",
    ]:
        _body(doc, text)

    _heading(doc, "8", "Frequently asked questions")
    for question, answer in FAQ:
        _sub(doc, question)
        _body(doc, answer, indent=0.2)

    _heading(doc, "9", "What this needs from us")
    _sub(doc, "Four decisions and two accounts")
    for i, text in enumerate([
        "**Publish the terms.** The signup page records which version an applicant accepted, and no version exists yet. Three clauses depend on the insurance position, the privacy policy, and our posture on training models with customer data - so it is really four documents with one on the critical path. Until it lands the front door cannot open.",
        "**Decide whether we take payment.** FAQ question one. It decides whether a whole workstream exists.",
        "**Confirm the rate card and the service levels for customer #1.** Billing supports more shapes than we need - a base per drop, per mile, per piece, per weight, with a minimum. What do we actually quote? And what did we promise per tier, given a miss now costs us a credit? The times we are running on are derived from our own operating constraints; the credit percentages are a placeholder and they are real money.",
        "**Provision Twilio and Rippling.** Text messaging carries driver login codes, customer tracking links, shop notifications and payment-dispute escalations. Payroll moves the money. Both are written and tested; neither has an account behind it.",
    ], start=1):
        _body(doc, f"{i}.  {text}", indent=0.2)

    doc.add_page_break()
    _heading(doc, "A", "Appendix - the screens, annotated")
    _body(doc, "Six screens, in the order the work moves through them: the page a prospect lands on, what "
               "we do with their application, the two screens a counter person lives in, the driver's stop, "
               "and what the person waiting for the part sees. Beside each one are the design decisions that "
               "are worth disagreeing with. **Every note is a choice we made and could unmake.**")
    _body(doc, "Each screen is transcribed rather than pictured - same fields, same order, same copy - "
               "because the wording and the sequence are what is being reviewed. What would help most is "
               "not whether they are pretty. It is where someone with a customer on the phone would stop "
               "to think.")

    counter = 0
    for screen in SCREENS:
        _sub(doc, screen["title"])
        _body(doc, screen["caption"])
        _screen(doc, screen["where"], screen["lines"])
        for claim, why in screen["notes"]:
            counter += 1
            _note(doc, counter, claim, why)
        _quote(doc, f"For discussion: {screen['ask']}")

    doc.add_page_break()
    _heading(doc, "B", "Appendix - capabilities at a glance")
    _body(doc, "For reference during the discussion rather than reading in order.")
    _table(doc, ["Capability", "What it means in practice"], CAPABILITIES)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}")
    print(f"  {out.stat().st_size // 1024} KB")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--out",
        type=Path,
        default=REPO / "docs" / "LMX_Link_Product_Design_Review.docx",
        help="where to write the .docx (default: docs/, which is gitignored)",
    )
    args = parser.parse_args()
    build(args.out.expanduser())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
