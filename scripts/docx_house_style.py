"""House style for generated Word documents (docs/DOCUMENT_STYLE.md).

Aptos, 0.5" margins, brand green `#0A6644`, the logo lockup on page one. Extracted
from `build_review_docx.py` when a second document needed the same treatment - the
alternative was two copies of the same twelve helpers drifting apart, which is exactly
the retrofitting the house style exists to prevent.

The public names here are the document vocabulary: `heading`, `body`, `sub`, `bullet`,
`table`, `quote`. Each takes the `Document` first so a build script reads as a sequence
of statements about the page rather than a tree of objects.

Two things python-docx has no API for - cell shading and paragraph borders - go in as
raw XML (`shade`, `rule`). That is not a workaround, it is the only route.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BRAND = RGBColor(0x0A, 0x66, 0x44)
BRAND_HEX = "0A6644"
TINT_HEX = "E6F1EB"
INK = RGBColor(0x14, 0x17, 0x1C)
INK_SOFT = RGBColor(0x3A, 0x41, 0x50)
INK_MUTED = RGBColor(0x5B, 0x64, 0x72)

REPO = Path(__file__).resolve().parent.parent
LOGO = REPO / "docs" / "LMX branding " / "lmx-logo-lockup-light.png"

FONT = "Aptos"
MONO = "Aptos Mono"

# `**bold**`, `*italic*` and `` `code` ``. Everything else is left alone: these
# documents are prose with the occasional file path in them, and a full markdown renderer
# here would be more surface than the output needs.
#
# Order matters. `**bold**` has to be tried before `*italic*` or the bold delimiters get
# eaten one asterisk at a time - and a lone asterisk in prose matches nothing, because
# both alternatives require a closing partner.
_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def shade(cell, hex_fill: str) -> None:
    """Table cell background. No python-docx API, so it goes in as XML."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def rule(paragraph, hex_color: str = BRAND_HEX, size: int = 12) -> None:
    """A coloured rule under a paragraph - the masthead and section heads."""
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def setup(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    style.paragraph_format.space_after = Pt(7)
    style.paragraph_format.line_spacing = 1.22
    # East-Asian font name too, or Word substitutes for any non-Latin run.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def logo(doc: Document) -> None:
    """The lockup at the top of page one, per house style - not a text label."""
    if not LOGO.exists():
        print(f"warning: logo not found at {LOGO} - continuing without it", file=sys.stderr)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(LOGO), height=Inches(0.30))


def kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = BRAND
    run.font.name = FONT


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = INK
    run.font.name = FONT


def byline(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = INK_MUTED
    run.font.name = FONT
    rule(p, BRAND_HEX, size=16)


def heading(doc: Document, number: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    n = p.add_run(f"{number}   ")
    n.font.size = Pt(10)
    n.font.bold = True
    n.font.color.rgb = BRAND
    n.font.name = FONT
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = INK
    run.font.name = FONT
    rule(p, "D8DCE2", size=6)


def _add_inline(paragraph, text: str, *, size: float, colour: RGBColor) -> None:
    """Fill a paragraph with runs, honouring **bold** and `mono`.

    A file path set in the body face is indistinguishable from prose, and these
    documents point at real paths that a reader is expected to open. The mono run is
    what makes `app/legal/content/terms.md` read as a thing rather than as a sentence.
    """
    for chunk in _INLINE.split(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk.strip("*`") if chunk[0] in "*`" else chunk)
        run.font.size = Pt(size)
        run.font.color.rgb = colour
        if chunk.startswith("**"):
            run.font.name = FONT
            run.font.bold = True
            run.font.color.rgb = INK
        elif chunk.startswith("`"):
            run.font.name = MONO
            run.font.size = Pt(size - 0.5)
            run.font.color.rgb = INK_SOFT
        elif chunk.startswith("*"):
            run.font.name = FONT
            run.font.italic = True
        else:
            run.font.name = FONT


def body(doc: Document, text: str, *, lead: bool = False, indent: float = 0.0) -> None:
    """One paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    _add_inline(p, text, size=11.5 if lead else 10.5, colour=INK_SOFT if lead else INK)


def bullet(doc: Document, text: str, *, indent: float = 0.25) -> None:
    """One bullet. A literal glyph with a hanging indent rather than Word's list
    styles, which carry numbering state across a document and are a poor bargain for a
    handful of flat lists."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent + 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    marker = p.add_run("•  ")
    marker.font.size = Pt(10.5)
    marker.font.color.rgb = BRAND
    marker.font.name = FONT
    _add_inline(p, text, size=10.5, colour=INK)


def quote(doc: Document, text: str) -> None:
    """A pull quote, marked with a left brand rule rather than italics - it is an
    assertion we want read, not an aside."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BRAND_HEX)
    bdr.append(left)
    p._p.get_or_add_pPr().append(bdr)
    _add_inline(p, text, size=11, colour=INK)


def sub(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = INK


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell, BRAND_HEX)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header.upper())
        run.font.name = FONT
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for index, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            if index % 2 == 1:
                shade(cells[i], TINT_HEX)
            cells[i].text = ""
            _add_inline(
                cells[i].paragraphs[0],
                f"**{value}**" if i == 0 and "**" not in value else value,
                size=9,
                colour=INK if i == 0 else INK_SOFT,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
